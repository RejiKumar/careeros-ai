"""Text extraction tests for PDF/DOCX/TXT uploads."""

from __future__ import annotations

import io

import pytest
from app.modules.resumes.text_extraction import MAX_EXTRACTED_CHARS, extract_text


def test_txt_extraction_strips_bom() -> None:
    content = b"\xef\xbb\xbfAda Lovelace\nLondon"
    assert extract_text("resume.txt", content) == "Ada Lovelace\nLondon"


def test_txt_extraction_truncates_long_files() -> None:
    content = ("x" * (MAX_EXTRACTED_CHARS + 500)).encode("utf-8")
    text = extract_text("resume.txt", content)
    assert len(text) == MAX_EXTRACTED_CHARS


def test_docx_extraction_reads_paragraphs_and_tables() -> None:
    from docx import Document

    document = Document()
    document.add_paragraph("Ada Lovelace")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Skill"
    table.rows[0].cells[1].text = "Python"
    buffer = io.BytesIO()
    document.save(buffer)

    text = extract_text("resume.docx", buffer.getvalue())

    assert "Ada Lovelace" in text
    assert "Skill" in text and "Python" in text


def test_corrupt_pdf_raises_value_error() -> None:
    with pytest.raises(ValueError):
        extract_text("resume.pdf", b"%PDF-1.4 not really a pdf")


def test_unknown_extension_raises_value_error() -> None:
    with pytest.raises(ValueError):
        extract_text("resume.exe", b"data")
