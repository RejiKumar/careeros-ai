"""Extract plain text from untrusted resume uploads (PDF/DOCX/TXT)."""

from __future__ import annotations

import io

MAX_EXTRACTED_CHARS = 100_000


def extract_text(filename: str, content: bytes) -> str:
    """Return up to MAX_EXTRACTED_CHARS characters of resume text.

    Extraction is best-effort; invalid files yield a friendly ValueError.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext == "pdf":
        return _pdf_text(content)
    if ext == "docx":
        return _docx_text(content)
    if ext == "txt":
        return _txt_text(content)
    raise ValueError("Unsupported file type. Use PDF, DOCX or TXT.")


def _pdf_text(content: bytes) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(content))
        pages = [page.extract_text() or "" for page in reader.pages]
        return _normalize("\n\n".join(pages))
    except Exception as exc:  # pypdf raises a variety of errors on corrupt files
        raise ValueError("Could not read the PDF file.") from exc


def _docx_text(content: bytes) -> str:
    from docx import Document

    try:
        document = Document(io.BytesIO(content))
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        return _normalize("\n".join(parts))
    except Exception as exc:
        raise ValueError("Could not read the DOCX file.") from exc


def _txt_text(content: bytes) -> str:
    return _normalize(content.decode("utf-8", errors="replace"))


def _normalize(text: str) -> str:
    text = text.lstrip("\ufeff")
    return text[:MAX_EXTRACTED_CHARS]
